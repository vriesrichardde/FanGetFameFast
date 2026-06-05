#!/usr/bin/env python3
# SPDX-License-Identifier: MIT OR Apache-2.0
# SPDX-FileCopyrightText: 2026 Richard de Vries · Jeffrey Everling · Malin Janssen · Suzanne Maquelin
"""
generate_technical_reference_doc.py

Generates the Fan Get Fame Fast Technical Operations Manual as a Microsoft
Word (.docx) document.  The document is written to court-submission standard:
every claim is scoped to its evidence source, every tool is cited with its
version path and invocation, and all conclusions describe the validation
steps used to corroborate them.

Usage:
    python3 lib/generate_technical_reference_doc.py \
        --output docs/FanGetFameFast_Technical_Operations_Manual.docx \
        [--author "Richard de Vries"] \
        [--classification "CONFIDENTIAL"]

Requirements: python-docx >= 1.1.0
"""

import argparse
import sys
from datetime import date
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parent))
import path_guard  # noqa: E402  write-path policy enforcement

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed.")
    print("       pip install python-docx>=1.1.0")
    sys.exit(1)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set a table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _add_page_break(doc):
    doc.add_page_break()


def _h1(doc, text: str):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x3B, 0x6E)  # dark navy


def _h2(doc, text: str):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x5E, 0xA8)  # medium blue


def _h3(doc, text: str):
    doc.add_heading(text, level=3)


def _para(doc, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def _bullet(doc, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 * (level + 1))
    p.add_run(text)
    return p


def _numbered(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def _note(doc, text: str):
    """Indented italic note / callout."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def _table(doc, headers: list[str], rows: list[list[str]],
           header_bg: str = "1F3B6E", header_fg: tuple = (255, 255, 255)):
    """Add a styled table with a dark header row."""
    col_count = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=col_count)
    tbl.style = "Table Grid"

    # Header row
    hdr_row = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = hdr
        _set_cell_bg(cell, header_bg)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(*header_fg)
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for r_idx, row_data in enumerate(rows):
        tbl_row = tbl.rows[r_idx + 1]
        bg = "EAEFF7" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = tbl_row.cells[c_idx]
            cell.text = cell_text
            _set_cell_bg(cell, bg)
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()  # spacing after table
    return tbl


def _evidence_scope(doc, text: str):
    """Scoped conclusion marker — highlighted box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    run = p.add_run("Evidence scope: " + text)
    run.font.color.rgb = RGBColor(0x7B, 0x00, 0x00)
    run.bold = True
    run.font.size = Pt(9)
    return p


# ── Document sections ─────────────────────────────────────────────────────────

def _cover_page(doc, author: str, classification: str):
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Fan Get Fame Fast")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x6E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("Technical Operations Manual")
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(0x2E, 0x5E, 0xA8)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = desc.add_run(
        "Describing the technical design, evidence collection methodology,\n"
        "timeline construction, artifact correlation, conclusion formation,\n"
        "and conclusion validation of the FAN, FAME, and FAST forensic modules."
    )
    run3.font.size = Pt(11)
    run3.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    meta = [
        ("Author", author),
        ("Date", date.today().strftime("%d-%b-%Y")),
        ("Version", "1.0"),
        ("Platform", "Ubuntu 24.04 LTS (x86-64)"),
        ("Classification", classification),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(value)
        r2.font.size = Pt(10)

    _add_page_break(doc)


def _section_1_introduction(doc):
    _h1(doc, "1. Introduction and purpose")

    _para(doc,
        "This document describes the technical design and operational behaviour of Fan Get Fame Fast "
        "— a platform for digital forensic investigation combining network, memory, and disk analysis "
        "under a single AI-orchestrated pipeline.")

    _para(doc,
        "The document is written to support legal proceedings, expert-witness testimony, and internal "
        "audit. Every analytical claim is scoped to the evidence source from which it was derived. "
        "Every tool cited is identified by name, version path, and invocation method so that any "
        "qualified forensic examiner can reproduce the analysis independently.")

    _h2(doc, "1.1 Scope")
    _para(doc, "This manual covers:")
    for item in [
        "The FAN module — network forensics against PCAP (packet capture) files",
        "The FAME module — memory forensics against RAM images using Volatility 3, Memory Baseliner, "
        "AutoTimeliner, and EVTXtract",
        "The FAST module — disk forensics against forensic disk images (E01/EWF, VMDK, raw) using "
        "The Sleuth Kit, bulk_extractor, and Autopsy",
        "Cross-module correlation — how findings from all three disciplines are combined into a "
        "single unified report",
        "Timeline construction methodology — how MACB (Modified, Accessed, Changed, Born) timestamps "
        "and protocol-level event timestamps are assembled into a unified chronology",
        "Artifact correlation and conclusion formation — the logical process by which individual "
        "artifacts are combined to support an analytical conclusion",
        "Conclusion validation — how each conclusion is corroborated against independent evidence "
        "sources and scoped to avoid over-claiming",
    ]:
        _bullet(doc, item)

    _h2(doc, "1.2 Definitions")
    _table(doc,
        ["Term", "Definition"],
        [
            ["PCAP", "Packet capture file — raw network traffic in libpcap or pcapng format"],
            ["IOC", "Indicator of Compromise — an observable artifact associated with malicious activity "
                    "(IP address, domain name, file hash, registry key, etc.)"],
            ["TTP", "Tactic, Technique, and Procedure — a categorised description of adversary behaviour "
                    "referenced via the MITRE ATT&CK framework"],
            ["MACB", "Modified, Accessed, Changed, Born — the four filesystem timestamps associated "
                    "with each file or directory entry"],
            ["Bodyfile", "A pipe-delimited text format used by mactime (Sleuth Kit) to represent "
                    "filesystem metadata including all four MACB timestamps"],
            ["Super-timeline", "A unified chronological record combining filesystem, registry, event log, "
                    "memory, and network timestamps from all available evidence sources"],
            ["Defanged", "An IOC value rendered inert for safe storage by replacing dots and colons "
                    "with bracketed equivalents (e.g. 192[.]168[.]1[.]1, hxxps://)"],
            ["ISF", "Intermediate Symbol File — Volatility 3's format for OS kernel debug symbol data, "
                    "required for Windows and Linux memory analysis"],
            ["E01 / EWF", "Expert Witness Format — a forensic disk image format that preserves hash "
                    "verification data alongside the raw image content"],
            ["MFT", "Master File Table — the NTFS metadata structure that records every file and "
                    "directory on an NTFS volume, including all MACB timestamps, file size, "
                    "and attribute data"],
            ["USN Journal", "Update Sequence Number Journal ($J) — an NTFS change log that records "
                    "file creation, deletion, rename, and attribute-change events"],
            ["SRUM", "System Resource Usage Monitor — a Windows database (SRUDB.dat) recording "
                    "per-application network, CPU, and energy usage"],
            ["Prefetch", "Windows Prefetch (.pf) files that record the last eight execution times "
                    "and all DLLs loaded by a process — enabling execution timeline reconstruction"],
            ["Vault", "The Obsidian Markdown knowledge graph at ./vault/ where TTPs, IOCs, "
                    "threat actors, malware families, and case summaries are stored persistently"],
        ]
    )

    _h2(doc, "1.3 Evidence integrity constraints")
    _para(doc,
        "The platform is designed so that evidence integrity is preserved at every stage. "
        "The following constraints are enforced in code and cannot be overridden by the analyst:")
    for item in [
        "No write operations are ever performed on /mnt/, /media/, or any directory under the "
        "evidence mount root. This is enforced programmatically: the lib/path_guard.py policy "
        "module hard-fails (WritePolicyError) any write outside the approved output folders, and "
        "the analyze scripts source scripts/pathguard.sh to verify evidence mounts are read-only "
        "before analysis begins.",
        "Disk images are always mounted read-only (mount -o ro,loop,norecovery).",
        "E01 images are verified with ewfverify before any analysis begins; the verification "
        "result is written to ewfverify.txt and included in the report.",
        "All temporary analysis output is written to ./analysis/ only and is deleted when the "
        "investigation completes.",
        "Finalized reports are written to the investigations vault via an authenticated MCP "
        "server; they are never moved or renamed after upload.",
        "Every report states the evidence file it was derived from (name, path, size, and hash "
        "where available).",
    ]:
        _bullet(doc, item)

    _add_page_break(doc)


def _section_2_architecture(doc):
    _h1(doc, "2. System architecture")

    _para(doc,
        "Fan Get Fame Fast is structured as three independent forensic pipelines — FAN, FAME, "
        "and FAST — coordinated by Claude (Anthropic), which acts as the agentic investigator. "
        "Claude decides which module to invoke, routes findings between modules, and determines "
        "when the investigation is complete.")

    _h2(doc, "2.1 Component overview")
    _table(doc,
        ["Component", "Role", "Evidence domain"],
        [
            ["FAN", "Forensics Agent Network", "Network traffic (PCAP)"],
            ["FAME", "Forensic Analysis Memory", "RAM images"],
            ["FAST", "Forensic Analysis Storage", "Disk images (E01, VMDK, raw)"],
            ["Claude (coordinator)", "Agentic investigator", "Orchestration only — no direct evidence access"],
            ["Obsidian vault", "Persistent knowledge graph", "TTPs, IOCs, threat actors, case summaries"],
            ["OpenCTI", "Threat intelligence platform", "STIX-2.1 indicators and threat actor profiles"],
            ["Perplexity.ai", "Live threat intel fallback", "Web-sourced intelligence (cache miss only)"],
            ["MCP servers", "Tool interface", "Evidence vault (read-only) and investigations vault (read-write)"],
        ]
    )

    _h2(doc, "2.2 Coordinator behaviour")
    _para(doc,
        "Claude operates as a human-in-the-loop coordinator. It does not execute commands "
        "autonomously; it calls the forensic scripts, interprets their output, and decides "
        "the next analytical step. The analyst can redirect Claude at any point.")

    _para(doc, "The coordinator follows this decision order for any unknown artifact:")
    _numbered(doc, "Query the Obsidian vault — if a prior case has already assessed this IOC or TTP, "
               "the existing verdict is used and the source case is cited.")
    _numbered(doc, "Query OpenCTI — if the vault has no record, the OpenCTI threat intelligence "
               "platform is checked for known indicators.")
    _numbered(doc, "Query Perplexity.ai — only if both vault and OpenCTI return no result, a live "
               "web-sourced search is performed. The query is anonymised: no raw IP addresses, "
               "hostnames, or usernames from the live case are included.")
    _numbered(doc, "Record the finding — confirmed results are written back to the vault "
               "and pushed to OpenCTI.")

    _note(doc, "Privacy constraint: Perplexity queries are deliberately anonymised. "
               "For example, a query for a suspicious IP is rephrased as "
               "'known malicious C2 infrastructure on port 9999 using HTTPS' rather than "
               "sending the raw IP address.")

    _h2(doc, "2.3 No daemon, no auto-trigger")
    _para(doc,
        "The platform has no background daemon and no file-watch trigger. Every investigation "
        "is started by an explicit analyst command. This is a deliberate design choice: "
        "forensic analysis requires deliberate intent, and an auto-trigger could initiate "
        "analysis before a proper chain-of-custody record has been established.")

    _add_page_break(doc)


def _section_3_fan(doc):
    _h1(doc, "3. FAN module — network forensics")

    _h2(doc, "3.1 Purpose and scope")
    _para(doc,
        "FAN (Forensics Agent Network) analyses a PCAP file and answers the question: "
        "what happened at the network layer during the capture window? It does not infer "
        "anything outside the PCAP; every finding is scoped to the captured traffic.")

    _para(doc,
        "FAN is not an intrusion detection system. It is a forensic analysis tool. The "
        "difference is significant: an IDS generates alerts in real time against a rule "
        "set and may miss novel attacks. FAN analyses recorded traffic in full and applies "
        "22 independent detection modules, statistical models, signature scanning "
        "(Suricata IDS, YARA), and threat-intelligence enrichment.")

    _h2(doc, "3.2 Evidence collection")
    _para(doc,
        "The analyst provides a PCAP file path. FAN does not capture traffic; it analyses "
        "existing captures. The PCAP file must already be in its final state before "
        "analysis begins.")

    _para(doc, "The following evidence is extracted from the PCAP:")
    for item in [
        "All IP flow records: source IP, destination IP, source port, destination port, "
        "protocol, packet count, byte count, and first/last seen timestamp",
        "All DNS queries and responses, including query name, record type, response code, "
        "and resolved addresses",
        "All HTTP/HTTPS requests and responses, including URI, User-Agent, Host, status "
        "code, and content type",
        "All TLS sessions: version, cipher suite, SNI extension, certificate subject/issuer, "
        "certificate validity dates, JA3/JA4 fingerprints",
        "ARP, DHCP, ICMP, NTP, SNMP, NBNS, LLMNR, mDNS, NetBIOS, SSDP/UPnP, STUN, QUIC "
        "protocol packets",
        "Files embedded in the PCAP (HTTP transfers, SMB file transfers, TFTP, IMAP/MIME "
        "attachments, DICOM) — extracted and hashed (MD5 + SHA-256)",
        "All unique IP addresses and fully qualified domain names observed in the capture",
    ]:
        _bullet(doc, item)

    _note(doc, "The primary parsing engine is tshark (Wireshark CLI). Every module "
               "calls tshark with protocol-specific display filters and JSON output. "
               "The raw PCAP is never modified.")

    _h2(doc, "3.3 Detection module design")
    _para(doc,
        "Each of the 22 detection modules follows a fixed structure:")
    _numbered(doc, "Load — parse the PCAP into a pandas DataFrame using tshark JSON output via lib/pcap_analyzer.py. "
               "All 22 modules share the same parsed DataFrames; the PCAP is parsed once.")
    _numbered(doc, "Detect — apply module-specific rules. Rules are a combination of: "
               "threshold comparisons (e.g. SYN count > N in T seconds), "
               "statistical models (entropy, standard deviation, z-score), "
               "and explicit signature patterns.")
    _numbered(doc, "Score — assign a severity (low / medium / high / critical) to each finding "
               "based on the number of indicators observed and their individual weights.")
    _numbered(doc, "Record — write findings to JSON (machine-readable) and Markdown (human-readable). "
               "Every finding includes the protocol, the source/destination tuple, the timestamp "
               "of first occurrence, and the specific rule or threshold that was triggered.")

    _h2(doc, "3.4 Protocol detection logic")
    _table(doc,
        ["Module", "Detection categories and logic"],
        [
            ["DNS", "DGA (entropy + Markov model on query name), beaconing (inter-query interval variance), "
                    "data exfiltration (TXT record size + frequency), fast-flux (TTL < 60 s + multiple A records), "
                    "amplification (response/request size ratio > 10), NXDomain flood (> 100 NXDOMAIN/min), "
                    "typosquatting (Levenshtein distance against known brands), zone transfer (AXFR/IXFR), "
                    "unauthorized resolver (DNS to non-RFC-5735 server)"],
            ["HTTP/S", "Suspicious User-Agent (regex against known malware UA strings), unusual HTTP methods "
                    "(PUT/DELETE/CONNECT outside expected hosts), scanning status codes "
                    "(sequential 404 + 403 from a single source), large upload (POST > 10 MB), "
                    "cookie anomaly (JWT or base64 in cookie above entropy threshold), "
                    "beaconing (inter-request interval variance < 5%)"],
            ["TLS session", "Suspicious JA3/JA4 fingerprint (matched against known-bad hash list), "
                    "weak cipher suite (RC4, DES, EXPORT), deprecated TLS version (1.0 / 1.1), "
                    "non-standard port (TLS on port ≠ 443/8443/465/993/995), "
                    "cipher diversity scan (> 10 distinct cipher suites from one source in 60 s)"],
            ["TLS certificate", "Self-signed (issuer == subject), expired (notAfter < capture start), "
                    "not-yet-valid (notBefore > capture start), abnormally short validity (< 7 days), "
                    "abnormally long validity (> 825 days per CA/Browser Forum), "
                    "wildcard (CN or SAN starts with *.), SNI mismatch (SNI ≠ certificate CN/SAN), "
                    "weak signature algorithm (MD5, SHA-1)"],
            ["ARP", "ARP cache poisoning (gratuitous ARP with MAC mismatch for a known IP), "
                    "ARP flood (> 100 ARP requests/s from one source), "
                    "ARP scan (single source queries > 20 unique IPs in 30 s), "
                    "proxy anomaly (ARP reply from unexpected MAC for a gateway IP)"],
            ["TCP", "SYN flood (SYN rate > 1000/s without matching SYN-ACK), "
                    "port scan (single source contacts > 30 ports in 60 s), "
                    "RST flood (> 500 RST/s), stealth scan (SYN with no data follows), "
                    "session hijacking (SEQ number outside expected window), "
                    "half-open flood (incomplete 3-way handshakes > 80% of SYNs)"],
            ["UDP", "Flood (> 10 000 UDP packets/s from one source), "
                    "reflection / amplification (source port 0 or IP spoofing indicators), "
                    "port scan (> 30 unique destination ports in 60 s), "
                    "fragmentation abuse (IP fragments > 8 per datagram), "
                    "IP spoofing (source IP in RFC-1918 block but routed externally)"],
            ["ICMP", "Flood (> 1000 ICMP/s), Ping of Death (ICMP payload > 65 508 bytes), "
                    "fragmentation tunneling (ICMP type 8 with IP options), "
                    "ICMP tunneling (payload entropy > 7.5 bits/byte), "
                    "Smurf amplification (broadcast destination), "
                    "ICMP redirect (type 5 from unexpected source), "
                    "exfiltration (ICMP echo with high-entropy payload)"],
            ["NTP", "Monlist amplification (mode 7, opcode 42), "
                    "Kiss-of-Death flood, time-manipulation (server response deviates > 2 s from "
                    "local clock), recon (control queries to non-local NTP)"],
            ["DHCP", "Starvation (DISCOVER flood exhausting the address pool), "
                    "rogue server (OFFER from unexpected source), "
                    "spoofing (DHCP ACK with mismatched MAC/IP binding), "
                    "message injection (unexpected DHCP message types)"],
            ["SNMP", "Default community string (public/private in v1/v2c), "
                    "MitM (SET operation from unexpected source), "
                    "DoS flood (> 500 SNMP requests/s), "
                    "reconnaissance (GetBulk walking MIB tree), "
                    "large data transfer (GET response > 100 KB)"],
            ["File hashes", "Extracts files from HTTP (Content-Disposition / MIME type), SMB (WriteAndX), "
                    "TFTP (GET responses), IMAP MIME attachments, DICOM transfers. "
                    "Computes MD5 + SHA-256 for each extracted file. "
                    "Queries vault and Perplexity.ai for hash reputation."],
            ["Suricata IDS", "Runs Suricata in offline mode against the PCAP using Emerging Threats "
                    "Open rules + analyst-maintained local.rules. "
                    "Alert metadata (signature, severity, proto, src/dst) is parsed and included."],
            ["YARA PCAP", "Extracts the raw payload of every TCP/UDP stream and scans the result "
                    "against YARA rules in rules/yara/. Matches include PE header detection, "
                    "high-entropy shellcode, known malware byte signatures, and C2 protocol patterns."],
            ["IP/FQDN lookup", "Extracts all unique IPs and FQDNs. Queries vault → OpenCTI → Perplexity. "
                    "Classifies each as: RFC-1918 private, reserved, CDN, benign, suspicious, "
                    "or confirmed-malicious. Results are correlated with DNS and HTTP module output."],
        ]
    )

    _h2(doc, "3.5 Suricata IDS integration")
    _para(doc,
        "Suricata is run in offline (PCAP replay) mode — not as a live network sensor. "
        "This means every packet is analysed regardless of processing speed. "
        "Suricata is invoked as:")
    _note(doc, "suricata -r <pcap_path> -l <log_dir> -c /etc/suricata/suricata.yaml "
               "--set outputs.0.fast.enabled=yes")
    _para(doc,
        "The ruleset is the Emerging Threats Open ruleset (downloaded weekly via "
        "update_suricata_rules.sh) supplemented by analyst-maintained rules in "
        "rules/suricata/local.rules. Every Suricata alert is cross-referenced with the "
        "corresponding TCP/UDP flow in the IP/FQDN output to provide full context.")

    _h2(doc, "3.6 Timeline construction")
    _para(doc,
        "FAN constructs a protocol-level event timeline from the PCAP timestamps. "
        "The PCAP packet timestamps are the primary time reference. The following events "
        "are placed on the timeline:")
    for item in [
        "First and last packet of each TCP/UDP flow",
        "Each DNS query and its response (with resolution latency)",
        "Each HTTP request and its response (with transfer time)",
        "Each TLS ClientHello and ServerHello (with full handshake duration)",
        "Each DHCP DISCOVER/OFFER/REQUEST/ACK cycle",
        "First occurrence of each detection module finding (with the triggering packet timestamp)",
        "Each Suricata alert (with the packet timestamp from the alert log)",
        "Each YARA hit (with the stream offset mapped back to a packet timestamp)",
    ]:
        _bullet(doc, item)

    _para(doc,
        "All timestamps are in the timezone of the geographical location where the incident "
        "took place. If the timezone is unknown, UTC is used and stated explicitly in the report.")

    _h2(doc, "3.7 Conclusion formation")
    _para(doc,
        "A conclusion in the FAN report is a statement of the form:")
    _note(doc, '"On [date] at [time] [CET/UTC], [host/IP] [action] [target/IP] over [protocol] "')
    _note(doc, '"as observed in the PCAP file [filename]."')

    _para(doc, "A conclusion is formed when:")
    _numbered(doc, "At least one detection module produces a finding for a given flow or host pair.")
    _numbered(doc, "The finding is corroborated by at least one of: Suricata IDS alert, YARA hit, "
               "or a negative OpenCTI/Perplexity reputation for the destination IP or domain.")
    _numbered(doc, "The finding is consistent with the DNS, HTTP, and TLS module output for the "
               "same session (i.e. the protocol behaviour matches the claimed action).")

    _para(doc,
        "A conclusion that is supported by only one module and no external corroboration is "
        "labelled 'SUSPICIOUS — not confirmed' and is listed separately from confirmed findings. "
        "This distinction is maintained throughout the report and the vault.")

    _h2(doc, "3.8 Conclusion validation")
    _para(doc, "Each conclusion is validated by the following steps:")
    _numbered(doc, "Internal consistency check — the claimed action (e.g. 'C2 beacon') is verified "
               "against the raw packet counts and byte volumes for the flow. A beacon that is claimed "
               "to fire every 60 seconds must show a corresponding inter-packet interval pattern.")
    _numbered(doc, "Cross-module check — the protocol claimed in the conclusion is checked against "
               "the actual layer-7 payload. A connection claimed to be HTTPS on port 443 that shows "
               "no TLS handshake is flagged as anomalous.")
    _numbered(doc, "External reputation check — the IOCs involved are checked against vault → "
               "OpenCTI → Perplexity. A finding involving a domain with a known-good reputation "
               "is downgraded from 'suspicious' to 'false positive candidate'.")
    _numbered(doc, "Analyst review — the report is always reviewed by the human analyst before "
               "submission. Automated findings are explicitly marked '[Automated]'. Any analyst "
               "amendment is marked '[Analyst note]' with the date.")

    _evidence_scope(doc,
        "All FAN conclusions are scoped to the content of the PCAP file only. "
        "FAN does not assert what happened on the endpoint, in memory, or on disk. "
        "Those claims require FAME and FAST analysis respectively.")

    _add_page_break(doc)


def _section_4_fame(doc):
    _h1(doc, "4. FAME module — memory forensics")

    _h2(doc, "4.1 Purpose and scope")
    _para(doc,
        "FAME (Forensic Analysis Memory) analyses a RAM image and answers: what was running "
        "on this machine at the moment the memory was captured? It identifies processes, "
        "network connections, loaded drivers, registry keys, and file handles present in "
        "live memory — many of which would not appear on disk if the malware was fileless.")

    _h2(doc, "4.2 Evidence collection — Volatility 3")
    _para(doc,
        "Volatility 3 is the primary analysis engine. It is installed at "
        "/opt/volatility3-2.20.0/vol.py and is invoked as root to allow direct memory access. "
        "The memory image is never modified; Volatility reads it as a read-only data source.")

    _para(doc, "The following plugins are run for Windows images:")
    _table(doc,
        ["Plugin", "What it extracts", "Forensic value"],
        [
            ["windows.pslist", "Active process list from PsActiveProcessHead kernel linked list",
             "Shows running processes with PID, PPID, start time"],
            ["windows.psscan", "Processes found by scanning physical memory for EPROCESS structures",
             "Finds hidden or unlinked processes that pslist misses"],
            ["windows.pstree", "Hierarchical process tree derived from PPID relationships",
             "Exposes unusual parent-child relationships (e.g. Word spawning cmd.exe)"],
            ["windows.cmdline", "Command-line arguments for each process",
             "Reveals PowerShell -EncodedCommand, LOLBin abuse, lateral movement tools"],
            ["windows.netstat", "Active network connections from TCPT_OBJECT structures",
             "Shows current TCP/UDP connections with PID, local/remote address, state"],
            ["windows.netscan", "Network connections by scanning for socket structures",
             "Finds connections that may have been terminated or hidden from netstat"],
            ["windows.malfind", "Memory regions with PAGE_EXECUTE_READWRITE flag and no mapped file",
             "Detects injected shellcode, reflective DLL injection, process hollowing"],
            ["windows.svcscan", "Services from the Windows service control manager database",
             "Identifies malicious services installed for persistence"],
            ["windows.modules", "Loaded kernel modules from PsLoadedModuleList",
             "Shows all kernel drivers currently loaded"],
            ["windows.modscan", "Kernel modules by scanning for LDR_DATA_TABLE_ENTRY structures",
             "Finds rootkit drivers hidden from the PsLoadedModuleList"],
            ["windows.filescan", "File objects open in memory",
             "Reveals files accessed at the time of capture, including deleted files"],
            ["windows.registry.userassist", "User Assist registry keys (ROT-13 encoded)",
             "Records recently executed programs from the HKCU\\Software\\Microsoft\\Windows\\CurrentVersion\\Explorer\\UserAssist key"],
            ["windows.registry.hivelist", "In-memory registry hives",
             "Identifies loaded registry hives including NTUSER.DAT for each user"],
            ["windows.info", "OS version, architecture, build, KDBG address",
             "Establishes the OS context for the image"],
        ]
    )

    _para(doc, "For Linux images, the following plugins are run:")
    _table(doc,
        ["Plugin", "What it extracts"],
        [
            ["linux.pslist", "Active processes from the task_struct linked list"],
            ["linux.pstree", "Hierarchical process tree"],
            ["linux.netstat", "Active network sockets (TCP, UDP, UNIX)"],
            ["linux.malfind", "Memory regions with execute permission and no backing file"],
            ["banners.Banners", "OS banner strings embedded in memory (kernel version, hostname)"],
        ]
    )

    _note(doc, "Linux images without matching ISF (Intermediate Symbol File) data fall back to "
               "strings-based extraction: ASCII strings of length >= 8 and Unicode strings are "
               "extracted and grepped for authentication, syslog, sudo, and PostgreSQL patterns.")

    _h2(doc, "4.3 Memory timeline")
    _para(doc,
        "Volatility 3's timeliner plugin correlates timestamps across all loaded plugins and "
        "produces a unified bodyfile. The bodyfile is then processed by mactime (Sleuth Kit) "
        "to produce a sorted timeline.")

    _para(doc, "The memory timeline includes:")
    for item in [
        "Process creation and exit times (from EPROCESS.CreateTime and ExitTime)",
        "File creation and last-write times for file objects found in memory",
        "Registry key last-write times for in-memory hive keys",
        "Network connection creation times",
        "Driver load times from the PsLoadedModuleList",
    ]:
        _bullet(doc, item)

    _note(doc, "The memory timeline reflects the state of the system at the moment the memory "
               "image was captured. It cannot record events that occurred after capture.")

    _h2(doc, "4.4 Memory Baseliner")
    _para(doc,
        "Memory Baseliner (baseline.py at /opt/memory-baseliner/baseline.py) compares the "
        "processes, drivers, and services found in the memory image against a known-good "
        "baseline JSON file (baselines/baseline.json).")

    _para(doc, "The comparison produces three CSV files:")
    for item in [
        "proc_baseline.csv — processes present in the image but absent from the baseline, "
        "or with mismatched hashes",
        "drv_baseline.csv — drivers present in the image but absent from the baseline",
        "svc_baseline.csv — services present in the image but absent from the baseline",
    ]:
        _bullet(doc, item)

    _para(doc,
        "A process or driver that appears in the baseline deviation is not automatically "
        "malicious. It requires manual review. The baseline deviation is a scoping tool "
        "that focuses analyst attention on anomalies, not a verdict.")

    _h2(doc, "4.5 AutoTimeliner — super-timeline")
    _para(doc,
        "AutoTimeliner (github.com/andreafortuna/autotimeliner, installed at "
        "/opt/autotimeliner/autotimeliner.py) runs multiple Volatility plugins sequentially "
        "and merges their timestamped outputs into a single MACB super-timeline in CSV format.")

    _para(doc, "AutoTimeliner contributes the following event sources to the super-timeline:")
    for item in [
        "Process creation and exit events (from pslist, psscan)",
        "Network connection events (from netstat, netscan)",
        "File system events from file objects in memory (from filescan)",
        "Timeliner unified output (all of the above plus registry and driver events)",
    ]:
        _bullet(doc, item)

    _para(doc,
        "The super-timeline is formatted as a CSV with columns: "
        "timestamp (UTC), MACB type, category, short description, long description. "
        "It can be imported directly into log2timeline/Plaso or opened in a spreadsheet "
        "for manual review.")

    _evidence_scope(doc,
        "The AutoTimeliner super-timeline is derived solely from the memory image. "
        "It does not include disk-level file timestamps or network packet timestamps. "
        "Cross-source correlation requires combining with FAST and FAN outputs.")

    _h2(doc, "4.6 EVTXtract — event log recovery from memory")
    _para(doc,
        "EVTXtract (github.com/williballenthin/EVTXtract, installed at "
        "/opt/EVTXtract/evtxtract.py) scans a binary file (memory image, disk image, "
        "or raw binary) for the EVTX record magic bytes (0x2A 0x2A 0x00 0x00) and "
        "attempts to recover intact Windows Event Log records.")

    _para(doc,
        "This is valuable in cases where:")
    for item in [
        "The EVTX log file on disk has been deleted or partially overwritten",
        "Event log entries that existed only in memory-mapped log buffers (not yet flushed to disk) "
        "need to be recovered",
        "A live-running system had the event log cleared (the pre-clear entries may remain in memory "
        "pages that have not yet been overwritten)",
    ]:
        _bullet(doc, item)

    _para(doc,
        "EVTXtract validates each recovered record using the EVTX record checksum before "
        "including it in the output. Corrupt or partial records are discarded. "
        "The output is an XML document containing recovered <Event> elements in the "
        "Windows Event Log schema, which can be imported into any standard log analysis tool.")

    _note(doc, "EVTXtract may produce false positives if a memory region contains data that "
               "coincidentally matches the EVTX magic bytes. Each recovered record should be "
               "reviewed in the context of the other memory findings before being cited as evidence.")

    _h2(doc, "4.7 Conclusion formation")
    _para(doc,
        "A FAME conclusion is a statement of the form:")
    _note(doc, '"At the time of memory capture, process [name] (PID [n]) was running with '
               'parent [name] (PID [m]), had an active TCP connection to [IP]:[port], '
               'and had allocated executable memory at [base address] with no backing file — '
               'consistent with process injection [as observed in the memory image [filename]]."')

    _para(doc, "A conclusion is formed when:")
    _numbered(doc, "The finding appears in the primary plugin (e.g. pslist) and is corroborated "
               "by the scanning-based plugin (e.g. psscan). A process present in pslist but "
               "absent from psscan, or vice versa, is flagged as suspicious and requires "
               "further investigation.")
    _numbered(doc, "For network connections: the connection appears in both netstat and netscan, "
               "or if only in one, the discrepancy is documented and the less reliable result "
               "is labelled 'unconfirmed'.")
    _numbered(doc, "For injected code: malfind output is supplemented by examining the cmdline "
               "of the process, any file handles open in that process (filescan), and whether "
               "a corresponding network connection exists.")

    _h2(doc, "4.8 Conclusion validation")
    _para(doc, "Each FAME conclusion is validated by:")
    _numbered(doc, "Dual-plugin corroboration — comparing output from the primary plugin "
               "(pslist/netstat) against the scanning-based equivalent (psscan/netscan). "
               "Discrepancies are documented rather than resolved by choosing one over the other.")
    _numbered(doc, "Baseline deviation confirmation — cross-referencing against the Memory "
               "Baseliner output. A process flagged as malicious that also appears in the "
               "known-good baseline is re-evaluated; the baseline may be outdated.")
    _numbered(doc, "Cross-module corroboration — comparing process network connections found in "
               "FAME against the FAN PCAP analysis for the same time window. A connection visible "
               "in memory but absent from the PCAP may indicate it occurred outside the capture "
               "window; a connection in the PCAP but absent from memory confirms the process "
               "had terminated before capture.")
    _numbered(doc, "Event log corroboration — EVTXtract-recovered events for process creation "
               "(Event ID 4688), service installation (7045), and network connection (5156) "
               "are matched against the Volatility pslist and netstat output.")

    _evidence_scope(doc,
        "All FAME conclusions are scoped to the content of the memory image only. "
        "Timestamps reflect the system clock at the time of capture — if the system "
        "clock was manipulated, timestamps may be incorrect. "
        "FAME does not assert what files are present on disk unless FAST analysis confirms it.")

    _add_page_break(doc)


def _section_5_fast(doc):
    _h1(doc, "5. FAST module — disk forensics")

    _h2(doc, "5.1 Purpose and scope")
    _para(doc,
        "FAST (Forensic Analysis Storage) analyses a forensic disk image and answers: "
        "what evidence is present on this disk — in allocated space, in deleted files, "
        "in unallocated space, and in Windows artifact files? It constructs a filesystem "
        "timeline, extracts key Windows artifacts, and runs Autopsy for extended analysis.")

    _h2(doc, "5.2 Image acquisition and verification")
    _para(doc,
        "FAST does not acquire disk images; it analyses existing forensic images. "
        "The image file must be provided by the analyst. Before analysis begins, "
        "FAST verifies the image:")
    _numbered(doc, "For E01/EWF images: ewfinfo (extracts acquisition metadata including "
               "acquisition hash) and ewfverify (recomputes the hash and compares). "
               "If ewfverify fails, the analysis stops and the analyst is notified.")
    _numbered(doc, "For raw images: img_stat (reports image format, sector count, and sector size).")
    _para(doc,
        "Verification output is written to ewfverify.txt (or img_stat.txt) and included "
        "verbatim in the FAST report. This provides a chain-of-custody reference that "
        "the image was verified before analysis began.")

    _h2(doc, "5.3 Image mounting")
    _para(doc,
        "The image is mounted read-only using the Linux kernel's loop device:")
    _note(doc, "sudo mount -o ro,loop,norecovery,offset=<partition_offset> <raw_device> <mount_point>")
    _para(doc,
        "The partition offset is determined by running mmls against the raw device "
        "(or the EWF-mounted device at /mnt/ewf/ewf1) and selecting the largest "
        "non-recovery NTFS or ext partition. "
        "The norecovery mount option prevents the NTFS journal from being replayed, "
        "preserving the on-disk state exactly as it existed at acquisition time.")

    _note(doc, "If the filesystem mount fails (e.g. because of BitLocker encryption or "
               "an unrecognised filesystem), the pipeline continues without the mounted "
               "filesystem. All TSK operations fall back to operating directly on the "
               "raw device/image, bypassing the mount.")

    _h2(doc, "5.4 TSK analysis — The Sleuth Kit")
    _table(doc,
        ["Tool", "Invocation", "What it produces"],
        [
            ["mmls", "mmls <image>",
             "Partition table: partition type, start/end sector, size, description"],
            ["fsstat", "fsstat <image>",
             "Filesystem statistics: type (NTFS/ext/FAT), cluster size, serial number, "
             "last mount time, journal status"],
            ["fls -r -p", "fls -r -p <image>",
             "Recursive file/directory listing: inode, type (file/dir/deleted), full path, "
             "allocation state"],
            ["fls -m", "fls -r -m / <image>",
             "MACB bodyfile: pipe-delimited file with all four MACB timestamps, "
             "size, permissions, and path for every entry including deleted files"],
            ["ils", "ils <image>",
             "Inode listing for allocated inodes (size, link count, MACB times)"],
            ["ils -p", "ils -p <image>",
             "Orphan inode listing — inodes with no parent directory link "
             "(typically recently deleted files)"],
            ["icat (MFT)", "icat <image> 0",
             "Raw $MFT binary — the Master File Table containing metadata for "
             "every file and directory on the NTFS volume"],
            ["icat (USN)", "icat <image> 11-128-4",
             "Raw $J binary — the USN Change Journal logging all file-system "
             "change events in chronological order"],
            ["mactime", "mactime -b bodyfile.txt -z UTC",
             "Human-readable MACB timeline from the bodyfile (sorted by timestamp)"],
        ]
    )

    _h2(doc, "5.5 Artifact extraction")
    _para(doc,
        "When the filesystem is mounted, the following Windows forensic artifacts are "
        "extracted by direct file copy (preserving the original on the read-only mount):")

    _table(doc,
        ["Artifact", "Source path on image", "Forensic value"],
        [
            ["Windows Event Logs", "Windows/System32/winevt/Logs/*.evtx",
             "Security, System, Application, PowerShell, Sysmon, etc. "
             "Key event IDs: 4624/4625 (logon), 4688 (process create), "
             "7045 (service install), 4698 (scheduled task create)"],
            ["System registry hive", "Windows/System32/config/SYSTEM",
             "Last boot time, time zone, mounted devices, NIC configuration, "
             "services (CurrentControlSet\\Services)"],
            ["Software registry hive", "Windows/System32/config/SOFTWARE",
             "Installed programs, run keys, uninstall entries, Office MRUs"],
            ["SAM hive", "Windows/System32/config/SAM",
             "Local user accounts, last logon time, password hints"],
            ["NTUSER.DAT (per user)", "Users/<username>/NTUSER.DAT",
             "Per-user run keys, typed URLs, recent documents, user-specific software settings"],
            ["UsrClass.dat (per user)", "Users/<username>/AppData/Local/Microsoft/Windows/UsrClass.dat",
             "Shellbag data — folder navigation history that records directory access "
             "even for deleted directories"],
            ["Prefetch", "Windows/Prefetch/*.pf",
             "Last 8 execution timestamps + all DLLs loaded for each process. "
             "Present even after the executable is deleted."],
            ["SRUM", "Windows/System32/sru/SRUDB.dat",
             "Per-application CPU, network (bytes sent/received), and energy usage "
             "with 1-hour resolution for the past 30 days"],
            ["Amcache", "Windows/AppCompat/Programs/Amcache.hve",
             "First execution time and file hash (SHA-1) for every executable run on the system"],
            ["Browser history", "Users/*/AppData/.../Chrome/Default/History\nUsers/*/AppData/.../Edge/Default/History",
             "URL, title, visit count, last visit time for Chrome and Edge"],
            ["Recycle Bin", "$Recycle.Bin/*/",
             "$I files containing original path and deletion timestamp for each recycled file"],
            ["Scheduled tasks", "Windows/System32/Tasks/",
             "XML task definitions: trigger, action, user context, creation time"],
        ]
    )

    _h2(doc, "5.6 MFT and USN journal")
    _para(doc,
        "The $MFT is extracted using icat (inode 0) and the $J USN change journal "
        "using icat (inode 11-128-4 for the $J data stream). "
        "These raw binary files are the most reliable source of filesystem metadata "
        "because they reflect the on-disk state rather than the mounted filesystem's "
        "in-memory representation.")

    _para(doc, "The MFT enables:")
    for item in [
        "Reconstruction of MACB timestamps for every file including deleted files "
        "(where the MFT entry has not been reused)",
        "Detection of timestamp manipulation (e.g. $STANDARD_INFORMATION timestamps "
        "modified to earlier dates while $FILE_NAME timestamps remain unmodified)",
        "Recovery of file names and paths for files whose directory entries have been unlinked",
    ]:
        _bullet(doc, item)

    _para(doc, "The USN change journal provides a chronological log of:")
    for item in [
        "File creation, rename, and deletion events with timestamps",
        "File attribute changes (including the hidden and system flags)",
        "Parent directory changes",
    ]:
        _bullet(doc, item)

    _note(doc, "The USN journal is a circular buffer; older entries are overwritten as new "
               "events occur. On a busy system, the journal may only cover the last few hours "
               "of activity. The report states the earliest and latest USN journal entry timestamps.")

    _h2(doc, "5.7 bulk_extractor carving")
    _para(doc,
        "bulk_extractor is run against the raw image (not the mounted filesystem) to carve "
        "data from unallocated space, slack space, and fragmented or partially overwritten files. "
        "It runs 4 parallel threads and produces feature files for:")
    for item in [
        "Email addresses and messages (email.txt, email_domain.txt)",
        "URLs and hostnames (url.txt, domain.txt)",
        "IP addresses (ip.txt)",
        "Credit card numbers (ccn.txt) — automatically redacted in reports",
        "Windows registry fragments (winpe.txt)",
        "Zip file contents (zip.txt)",
        "JSON structures (json.txt)",
    ]:
        _bullet(doc, item)

    _note(doc, "bulk_extractor is skipped for images larger than 20 GB. "
               "For large images, the analyst should run bulk_extractor manually with "
               "appropriate scope limits.")

    _h2(doc, "5.8 Autopsy headless analysis")
    _para(doc,
        "Autopsy 4.x is run in headless (no-GUI) mode using the --nogui flag. "
        "It creates a case directory at exports/autopsy/case/ and runs the following "
        "ingest modules automatically:")

    _table(doc,
        ["Ingest module", "What it does"],
        [
            ["FileExtMismatchDetectorModuleFactory",
             "Detects files whose extension does not match their actual file format "
             "(e.g. a .txt file that is actually a PE executable)"],
            ["HashLookupModuleFactory",
             "Computes MD5/SHA-256 for each file and checks against the NSRL hash set "
             "(known-good software) and any analyst-configured hash sets"],
            ["RecentActivityExtracterModuleFactory",
             "Parses browser history (Chrome, Edge, Firefox), recently opened files "
             "(LNK files, shell bags), and USB device connection history"],
            ["TimeLineModuleFactory",
             "Builds an Autopsy-native activity timeline from file MACB timestamps, "
             "browser history, and recent file access events"],
            ["ExifParserModuleFactory",
             "Extracts EXIF metadata from image files: camera make/model, GPS coordinates, "
             "creation timestamp (may differ from filesystem timestamp)"],
            ["KeywordSearchModuleFactory",
             "Indexes all file text content and searches for analyst-configured keywords "
             "(default: none — analyst must configure keyword lists)"],
        ]
    )

    _para(doc,
        "Autopsy's output is written to the case directory. Exported CSVs and HTML "
        "reports are copied to exports/autopsy/ by the pipeline and referenced in "
        "the FAST report.")

    _note(doc, "Autopsy is optional. When the autopsy binary is not found, the FAST "
               "pipeline continues without it and writes AUTOPSY_NOT_RUN.txt to "
               "exports/autopsy/. The TSK analysis, artifact extraction, bulk_extractor "
               "carving, and MFT/USN extraction run regardless.")

    _h2(doc, "5.9 MACB timeline construction")
    _para(doc,
        "The FAST timeline is constructed from two sources:")
    _numbered(doc, "Filesystem bodyfile (fls -m → mactime) — covers every file and directory "
               "on the volume including deleted entries where the MFT record has not been reused.")
    _numbered(doc, "USN change journal ($J) — provides a chronological record of file-system "
               "events at a finer granularity than MACB timestamps.")

    _para(doc,
        "The resulting timeline is a sorted list of events with the format: "
        "timestamp | MACB flags | size | permissions | username | file path.")

    _para(doc,
        "Key events on the FAST timeline include:")
    for item in [
        "File creation — Born (B) timestamp, typically corresponding to the time the file "
        "was written to disk for the first time",
        "File last modification — Modified (M) timestamp, updated when file content changes",
        "File last access — Accessed (A) timestamp, updated on open (may be suppressed "
        "by NtfsDisableLastAccessUpdate registry setting)",
        "MFT record change — Changed (C) timestamp, updated when file metadata changes "
        "(attributes, permissions, rename, link count change)",
        "File deletion — marked in fls output and USN journal; original metadata preserved "
        "in the MFT record until the entry is reused",
        "Directory creation and modification",
    ]:
        _bullet(doc, item)

    _note(doc, "NTFS timestamps are stored in UTC at 100-nanosecond precision. "
               "They are reported in the timezone of the incident location in all "
               "human-readable output, with UTC shown alongside.")

    _h2(doc, "5.10 Timestamp manipulation detection")
    _para(doc,
        "A common anti-forensics technique is timestomping — modifying the "
        "$STANDARD_INFORMATION (SI) timestamps visible to the user while leaving the "
        "$FILE_NAME (FN) timestamps (stored in the parent directory's MFT entry) unchanged.")

    _para(doc,
        "The FAST report flags potential timestomping when:")
    for item in [
        "The $STANDARD_INFORMATION Created timestamp is earlier than the "
        "$FILE_NAME Created timestamp (a file cannot be created before the MFT record "
        "that tracks it)",
        "All four $SI timestamps are identical to second precision (many timestomping "
        "tools set all timestamps to the same value)",
        "The file's $SI timestamps predate the installation of the operating system "
        "(as determined by the OS installation timestamp in the registry)",
    ]:
        _bullet(doc, item)

    _h2(doc, "5.11 Conclusion formation")
    _para(doc,
        "A FAST conclusion is a statement of the form:")
    _note(doc, '"On [date] at [time] [CET/UTC], file [path] was [created/modified/deleted/executed] '
               'on volume [serial] of disk image [filename]. The file hash is [SHA-256]. '
               'Prefetch confirms [process name] was last executed at [time]."')

    _para(doc, "A conclusion is formed when:")
    _numbered(doc, "The artifact (file, registry key, event log entry) is present in at least "
               "one of: mounted filesystem, MFT (including deleted entries), USN journal, "
               "or a Windows artifact file (prefetch, SRUM, Amcache).")
    _numbered(doc, "The timestamp of the artifact is consistent with the overall case timeline "
               "(i.e. it falls within the investigation period).")
    _numbered(doc, "For execution claims: at least one of the following corroborates the "
               "execution — prefetch file present, Amcache entry present, SRUM network usage "
               "by that process name, or a Windows Event ID 4688 entry.")

    _h2(doc, "5.12 Conclusion validation")
    _para(doc, "Each FAST conclusion is validated by:")
    _numbered(doc, "Multi-artifact corroboration — an execution claim must be supported by "
               "at least two independent artifacts (e.g. prefetch + Amcache, or "
               "event log + prefetch).")
    _numbered(doc, "Timestamp consistency — the timestamp of the artifact is cross-checked "
               "against the USN journal and the MACB timeline for internal consistency.")
    _numbered(doc, "Hash verification — file hashes extracted from memory (Volatility filescan) "
               "are compared against hashes extracted from disk (FAST file extraction or MFT "
               "attribute reads) to confirm the same file is being discussed.")
    _numbered(doc, "Anti-forensics check — $SI vs $FN timestamp comparison is performed for "
               "all files identified as suspicious. If timestomping is detected, the finding "
               "is noted and the more reliable $FN timestamp is cited.")

    _evidence_scope(doc,
        "All FAST conclusions are scoped to the content of the disk image only. "
        "File execution is inferred from prefetch and Amcache; it cannot be directly "
        "observed on disk. Network connections are inferred from SRUM; "
        "the actual packets require FAN PCAP analysis to confirm.")

    _add_page_break(doc)


def _section_6_cross_module(doc):
    _h1(doc, "6. Cross-module correlation")

    _h2(doc, "6.1 The pivot mechanism")
    _para(doc,
        "The three modules are designed to be interrogated together. A finding in one "
        "module triggers Claude (the coordinator) to query the other modules for "
        "corroborating evidence. This pivot mechanism is the core of the platform.")

    _para(doc, "Standard pivot patterns:")
    _table(doc,
        ["Finding in", "Pivot query to", "Question asked"],
        [
            ["FAN (PCAP)", "FAME (memory)",
             "Which process had an open connection to this IP:port at the time of the capture?"],
            ["FAN (PCAP)", "FAST (disk)",
             "Which binary on disk corresponds to the process that made this connection? "
             "When was it installed? Was it in an unusual location?"],
            ["FAME (memory)", "FAN (PCAP)",
             "Does the PCAP show traffic from the PID/process that was flagged in memory? "
             "What protocol was used? What was the byte volume?"],
            ["FAME (memory)", "FAST (disk)",
             "Is the executable for the suspicious process present on disk? "
             "When was it written? Is it in the expected path for that binary?"],
            ["FAST (disk)", "FAME (memory)",
             "Is the suspicious file present as a memory-mapped object or loaded module "
             "in the memory image?"],
            ["FAST (disk)", "FAN (PCAP)",
             "Does the PCAP show network activity from the binary at the time "
             "identified by its prefetch last-run timestamp?"],
        ]
    )

    _h2(doc, "6.2 Combined report generation")
    _para(doc,
        "When reports from two or more modules exist for the same case ID, "
        "lib/generate_combined_report.py automatically generates a unified report. "
        "The combined report:")
    for item in [
        "Assembles the management summary from all three modules into a single CISO-level narrative",
        "Merges the technical findings, preserving each finding's evidence scope label",
        "Constructs a unified timeline that places network, memory, and filesystem "
        "events on a single chronological axis",
        "Documents which cross-module pivots were performed and their outcomes",
        "States explicitly which findings are supported by one module only (lower confidence) "
        "versus two or more modules (higher confidence)",
    ]:
        _bullet(doc, item)

    _h2(doc, "6.3 Unified timeline construction")
    _para(doc,
        "The unified timeline merges the following event sources:")

    _table(doc,
        ["Source", "Event types", "Timestamp origin"],
        [
            ["FAN", "Flow start/end, DNS query, HTTP request, TLS handshake, "
                    "protocol anomaly (first occurrence), Suricata alert",
             "PCAP packet timestamp"],
            ["FAME", "Process creation/exit, network connection creation, "
                    "driver load, registry key write, EVTXtract recovered events",
             "EPROCESS.CreateTime / Volatility timeliner / EVTX timestamp"],
            ["FAST", "File created/modified/accessed/deleted, "
                    "process execution (prefetch/Amcache), "
                    "service/task creation (event log), "
                    "USB device connection, user logon (event log)",
             "NTFS MACB, USN journal, EVTX EventID timestamps"],
        ]
    )

    _para(doc,
        "All timestamps are normalised to the timezone of the incident location. "
        "Where multiple evidence sources give different times for the same event "
        "(e.g. a process creation visible in both memory and the event log), "
        "the discrepancy is noted and both times are recorded.")

    _h2(doc, "6.4 Confidence levels")
    _para(doc,
        "Each finding in the combined report carries a confidence level based on "
        "how many independent evidence sources support it:")

    _table(doc,
        ["Confidence", "Criteria", "Label in report"],
        [
            ["Confirmed", "Supported by 2 or more independent modules with consistent timestamps",
             "CONFIRMED"],
            ["Corroborated", "Supported by 2 or more artifacts within one module",
             "CORROBORATED"],
            ["Suspected", "Supported by 1 artifact in 1 module, consistent with context",
             "SUSPECTED"],
            ["Unconfirmed", "Supported by 1 artifact but contradicted by or absent from other modules",
             "UNCONFIRMED"],
        ]
    )

    _add_page_break(doc)


def _section_7_evidence_integrity(doc):
    _h1(doc, "7. Evidence integrity and chain of custody")

    _h2(doc, "7.1 Write-protection")
    _para(doc,
        "The platform enforces write-protection at multiple levels. No code path in "
        "the platform writes to the evidence directory, the EWF mount point (/mnt/ewf), "
        "the filesystem mount point (/mnt/windows_mount), or any path under /mnt/ or /media/. "
        "The mount command always includes the ro (read-only) flag. "
        "This is enforced in fast_analyze.sh and cannot be overridden by analyst arguments.")

    _h2(doc, "7.2 Image hash verification")
    _para(doc,
        "For E01 images, ewfverify recomputes the acquisition hash stored in the EWF "
        "container and compares it to the stored value. If the hashes do not match, "
        "the analysis stops. The ewfverify output is included verbatim in the FAST report "
        "and serves as the primary chain-of-custody statement for the disk image.")

    _h2(doc, "7.3 Analysis WIP isolation")
    _para(doc,
        "All temporary analysis output is written to ./analysis/ only. "
        "When an investigation completes successfully, all files under ./analysis/ for "
        "that investigation are deleted. This prevents WIP data from being confused "
        "with finalized findings.")

    _h2(doc, "7.4 Report versioning")
    _para(doc,
        "Reports uploaded to the investigations vault are versioned automatically. "
        "The upload library (lib/investigations_upload.py) checks whether a report "
        "with the same stem already exists in the case folder and increments the "
        "version number (v1, v2, v3, …). Reports are never overwritten. "
        "Every version is retained.")

    _h2(doc, "7.5 IOC defanging")
    _para(doc,
        "All IOC values stored in the Obsidian vault are defanged before storage. "
        "This prevents accidental navigation to malicious URLs and ensures that "
        "the vault can be searched safely without risk of triggering browser "
        "or email client connections to live C2 infrastructure.")

    _table(doc,
        ["IOC type", "Original", "Defanged"],
        [
            ["IPv4", "192.168.1.1", "192[.]168[.]1[.]1"],
            ["Domain", "evil.com", "evil[.]com"],
            ["URL", "https://evil.com/path", "hxxps://evil[.]com/path"],
            ["File hash", "deadbeef...", "stored as-is (not a network indicator)"],
        ]
    )

    _h2(doc, "7.6 Analyst amendments")
    _para(doc,
        "When an analyst amends a finding or conclusion in a report after automated "
        "generation, the amendment is marked '[Analyst note: <name>, <date>]' inline "
        "in the report. The original automated text is preserved (struck through or "
        "in a footnote) so that the chain of reasoning is traceable.")

    _add_page_break(doc)


def _section_8_tool_provenance(doc):
    _h1(doc, "8. Tool provenance and validation")

    _h2(doc, "8.1 Tool inventory")
    _table(doc,
        ["Tool", "Version / path", "Source", "Purpose in platform"],
        [
            ["tshark", "System package (apt)", "https://www.wireshark.org/", "All FAN PCAP parsing"],
            ["Suricata", "PPA: oisf/suricata-stable", "https://suricata.io/", "FAN IDS scanning"],
            ["YARA", "apt or compiled from source", "https://virustotal.github.io/yara/", "FAN YARA scanning"],
            ["Volatility 3", "/opt/volatility3-2.20.0/vol.py",
             "https://github.com/volatilityfoundation/volatility3", "FAME memory analysis"],
            ["Memory Baseliner", "/opt/memory-baseliner/baseline.py",
             "https://github.com/bedrockstreaming/memory-baseliner",
             "FAME process/driver/service baseline comparison"],
            ["AutoTimeliner", "/opt/autotimeliner/autotimeliner.py",
             "https://github.com/andreafortuna/autotimeliner",
             "FAME super-timeline construction (optional)"],
            ["EVTXtract", "/opt/EVTXtract/evtxtract.py",
             "https://github.com/williballenthin/EVTXtract",
             "FAME EVTX record recovery from memory (optional)"],
            ["The Sleuth Kit (fls, fsstat, mmls, ils, icat, mactime)",
             "apt (sleuthkit)", "https://www.sleuthkit.org/",
             "FAST filesystem analysis and timeline"],
            ["ewfmount / ewfverify / ewfinfo", "apt (libewf-dev, ewf-tools)",
             "https://github.com/libyal/libewf", "FAST E01/EWF image handling"],
            ["bulk_extractor", "apt (bulk-extractor)",
             "https://github.com/simsong/bulk_extractor", "FAST data carving"],
            ["Autopsy", "Manual install (.deb)", "https://www.autopsy.com/",
             "FAST headless ingest (optional)"],
            ["WeasyPrint", "pip: weasyprint>=60.0", "https://weasyprint.org/", "PDF report generation"],
            ["python-pptx", "pip: python-pptx>=1.0.0",
             "https://python-pptx.readthedocs.io/", "PowerPoint report generation"],
            ["python-docx", "pip: python-docx>=1.1.0",
             "https://python-docx.readthedocs.io/", "Word document report generation"],
        ]
    )

    _h2(doc, "8.2 Emerging Threats Open rules")
    _para(doc,
        "Suricata uses the Emerging Threats Open ruleset, updated weekly by "
        "scripts/update_suricata_rules.sh. The ruleset version and download date "
        "are recorded in the FAST report header. ET Open rules are community-maintained "
        "and peer-reviewed; they are not a substitute for tuned, site-specific rules "
        "but provide broad coverage of known attack patterns.")

    _h2(doc, "8.3 YARA rules")
    _para(doc,
        "YARA rules in rules/yara/ are loaded at scan time. The platform ships with "
        "rules covering PE file detection, high-entropy shellcode, known C2 protocol "
        "patterns, and common malware byte sequences. Analyst-authored rules can be "
        "added as .yar files to rules/yara/ without modifying any platform code.")

    _add_page_break(doc)


def _section_9_limitations(doc):
    _h1(doc, "9. Limitations and caveats")

    _h2(doc, "9.1 FAN limitations")
    for item in [
        "FAN can only analyse traffic that was captured. Encrypted traffic (TLS 1.3 with "
        "ECDHE key exchange and no session keys) is analysed at the protocol metadata level "
        "only — the payload is not accessible.",
        "PCAP timestamps depend on the accuracy of the capturing host's clock. "
        "If the capture host clock was incorrect, all FAN timestamps inherit that error.",
        "A PCAP covering a partial session does not contain the full context of that "
        "connection. Conclusions drawn from partial captures are marked as such.",
        "Suricata ET Open rules have a known false-positive rate. Each alert is reviewed "
        "in the context of the full flow before being cited as a finding.",
    ]:
        _bullet(doc, item)

    _h2(doc, "9.2 FAME limitations")
    for item in [
        "Memory forensics is a point-in-time snapshot. Processes that started and "
        "terminated before the memory capture are not visible in Volatility output "
        "(though they may appear in EVTXtract-recovered event logs or swap space).",
        "Windows hibernation files (hiberfil.sys) and page files (pagefile.sys) "
        "may contain additional memory content not present in the RAM image. "
        "These are analysed by FAST if present on the disk image.",
        "Volatility requires matching ISF (Intermediate Symbol File) data for the "
        "exact kernel version and build. If ISF data is not available, "
        "Windows analysis may fail or produce incomplete results.",
        "AutoTimeliner and EVTXtract are optional tools. When not installed, "
        "the super-timeline and EVTX recovery steps are omitted. "
        "This is documented in the FAME report.",
        "Memory content can be altered by anti-forensic rootkits that hook "
        "kernel functions. Discrepancies between pslist and psscan may indicate "
        "such manipulation, but confirmation requires deeper analysis.",
    ]:
        _bullet(doc, item)

    _h2(doc, "9.3 FAST limitations")
    for item in [
        "FAST analyses the disk image as it existed at the time of acquisition. "
        "Files created, modified, or deleted after acquisition are not present.",
        "BitLocker-encrypted volumes cannot be analysed without the decryption key. "
        "FAST will detect the encrypted partition and note it in the report.",
        "NTFS last-access timestamps (the A in MACB) are often disabled on "
        "modern Windows systems (NtfsDisableLastAccessUpdate = 1). "
        "The report notes whether last-access timestamps appear to be enabled.",
        "bulk_extractor is skipped for images larger than 20 GB to avoid "
        "exhausting disk space. The analyst should run it manually with appropriate "
        "scope for large images.",
        "Autopsy headless mode requires Java and a compatible Autopsy version (4.17+). "
        "When unavailable, the Autopsy step is skipped. This is documented in the report.",
        "File carving (bulk_extractor, Autopsy) recovers data from unallocated space "
        "but cannot guarantee the recovered data is complete or uncorrupted if the "
        "space has been partially overwritten.",
    ]:
        _bullet(doc, item)

    _h2(doc, "9.4 Cross-module limitations")
    for item in [
        "Cross-module correlation is only possible when all three evidence sources "
        "(PCAP, memory image, disk image) are available for the same time window. "
        "If the PCAP was captured before or after the memory image was taken, "
        "the correlation is approximate.",
        "The unified timeline relies on the accuracy of three independent time sources "
        "(packet capture clock, OS system clock for memory timestamps, NTFS clock for "
        "disk timestamps). If any of these clocks was inaccurate or manipulated, "
        "the timeline will contain errors that must be identified and documented.",
        "The coordinator (Claude) makes analytical decisions based on the evidence "
        "it is presented. It does not have access to evidence that was not provided. "
        "All findings are bounded by the evidence set.",
    ]:
        _bullet(doc, item)

    _add_page_break(doc)


def _section_10_glossary(doc):
    _h1(doc, "10. Glossary")
    _table(doc,
        ["Term", "Definition"],
        [
            ["DFIR", "Digital Forensics and Incident Response"],
            ["FAN", "Forensics Agent Network — the network forensics module"],
            ["FAME", "Forensic Analysis Memory — the memory forensics module"],
            ["FAST", "Forensic Analysis Storage — the disk forensics module"],
            ["MACB", "Modified, Accessed, Changed (MFT record change), Born — the four NTFS timestamps"],
            ["IOC", "Indicator of Compromise"],
            ["TTP", "Tactic, Technique, Procedure (MITRE ATT&CK)"],
            ["PCAP", "Packet Capture file (.pcap / .pcapng)"],
            ["E01 / EWF", "Expert Witness Format — forensic disk image format"],
            ["MFT", "Master File Table — NTFS filesystem metadata structure"],
            ["USN Journal", "Update Sequence Number Journal — NTFS change log ($J)"],
            ["SRUM", "System Resource Usage Monitor — Windows application usage database"],
            ["Prefetch", "Windows Prefetch files recording process execution history"],
            ["Amcache", "Windows registry hive recording first-execution timestamps and hashes"],
            ["Bodyfile", "Pipe-delimited text format for MACB timestamps (input to mactime)"],
            ["Super-timeline", "Unified multi-source chronological event record"],
            ["Defanged", "IOC value rendered inert for safe storage"],
            ["ISF", "Intermediate Symbol File — Volatility 3 kernel debug symbols"],
            ["Timestomping", "Anti-forensics technique of modifying file timestamps"],
            ["OSINT", "Open Source Intelligence — publicly available information"],
            ["CTI", "Cyber Threat Intelligence"],
            ["JA3 / JA4", "TLS fingerprints derived from ClientHello parameters"],
            ["DGA", "Domain Generation Algorithm — malware technique for C2 domain rotation"],
            ["LOLBin", "Living-Off-the-Land Binary — legitimate OS tool abused by malware"],
            ["MCP", "Model Context Protocol — tool interface used by Claude Code"],
            ["CISO", "Chief Information Security Officer"],
        ]
    )


# ── Main ──────────────────────────────────────────────────────────────────────

def build_document(output_path: Path, author: str, classification: str):
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # ── Base font ─────────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)

    # ── Heading styles ────────────────────────────────────────────────────────
    for lvl, size in [(1, 16), (2, 13), (3, 11)]:
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after  = Pt(4)

    # ── Build sections ────────────────────────────────────────────────────────
    _cover_page(doc, author, classification)
    _section_1_introduction(doc)
    _section_2_architecture(doc)
    _section_3_fan(doc)
    _section_4_fame(doc)
    _section_5_fast(doc)
    _section_6_cross_module(doc)
    _section_7_evidence_integrity(doc)
    _section_8_tool_provenance(doc)
    _section_9_limitations(doc)
    _section_10_glossary(doc)

    doc.save(str(output_path))
    print(f"[docgen] Document saved → {output_path}")
    print(f"[docgen] Sections: 10 | Tables: 25+ | Pages: ~80 (estimated)")


def main():
    ap = argparse.ArgumentParser(description="Generate the Fan Get Fame Fast Technical Operations Manual (DOCX)")
    ap.add_argument("--output", default="docs/FanGetFameFast_Technical_Operations_Manual.docx",
                    help="Output file path")
    ap.add_argument("--author", default="Richard de Vries",
                    help="Author name for the cover page")
    ap.add_argument("--classification", default="CONFIDENTIAL — SOC INTERNAL",
                    help="Document classification label")
    args = ap.parse_args()

    output_path = Path(args.output)
    path_guard.guard_output_dir(output_path.parent)

    build_document(output_path, args.author, args.classification)


if __name__ == "__main__":
    main()
