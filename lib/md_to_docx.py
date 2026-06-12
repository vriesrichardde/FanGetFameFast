#!/usr/bin/env python3
# SPDX-License-Identifier: MIT OR Apache-2.0
# SPDX-FileCopyrightText: 2026 Richard de Vries · Jeffrey Everling · Malin Janssen · Suzanne Maquelin · Joost Beekman
"""
md_to_docx.py — Generic Markdown -> Word (.docx) renderer.

Renders a hand-authored campaign report markdown file as a DOCX: the
document's H1 becomes a cover heading with a metadata table (parsed from the
header table that follows it), `##`/`###` headings become Word headings,
pipe tables become real Word tables, bullet/numbered lists become Word list
paragraphs, and `> Claude:` blockquote authoring directives are dropped —
they must never reach the rendered document.

Usage (CLI):
    python3 lib/md_to_docx.py /path/to/campaign_report.md \\
        --output /path/to/campaign_report.docx \\
        --case-id CASE-2026-001

Python API:
    from lib.md_to_docx import convert
    docx_path = convert(md_path, output_path, case_id="CASE-2026-001")
"""
from __future__ import annotations

import argparse
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import path_guard  # noqa: E402  write-path policy enforcement

_TABLE_ROW_RE = re.compile(r"^\|(.+)\|\s*$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\d+[.)]\s+(.*)$")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def _xml_safe(text: str) -> str:
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", str(text))


def _strip_md(text: str) -> str:
    text = re.sub(r"\[([^\]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    return text.strip()


def _is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{2,}:?", c) for c in cells)


def _parse_table_rows(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        m = _TABLE_ROW_RE.match(line.strip())
        if not m:
            continue
        cells = [c.strip() for c in m.group(1).split("|")]
        if _is_separator_row(cells):
            continue
        rows.append(cells)
    return rows


def convert(md_path: Path, output_path: Path, case_id: str = "",
            title: str = "", date_str: str = "") -> Path:
    """Render *md_path* as a DOCX. Returns the output path."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise SystemExit(
            "[md_to_docx] 'python-docx' package not found.\n"
            "Install: pip3 install python-docx"
        )

    md_path = Path(md_path)
    if not md_path.exists():
        raise FileNotFoundError(f"Markdown file not found: {md_path}")

    out = path_guard.assert_writable(Path(output_path))
    path_guard.guard_output_dir(out.parent)

    if not date_str:
        date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    def _h(text, level):
        p = doc.add_heading(_xml_safe(text), level=level)
        if p.runs:
            p.runs[0].font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    def _p(text, bold=False, italic=False, style=None):
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        run = p.add_run(_xml_safe(text))
        run.bold = bold
        run.italic = italic

    def _note(text):
        p = doc.add_paragraph()
        run = p.add_run(_xml_safe(text))
        run.italic = True
        run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    def _table(rows: list[list[str]]):
        if not rows:
            return
        header, *data = rows
        n_cols = len(header)
        tbl = doc.add_table(rows=1 + len(data), cols=n_cols)
        tbl.style = "Table Grid"
        for j, h in enumerate(header):
            tbl.rows[0].cells[j].text = _xml_safe(_strip_md(h))
            for run in tbl.rows[0].cells[j].paragraphs[0].runs:
                run.bold = True
        for i, row in enumerate(data, start=1):
            for j in range(n_cols):
                val = _strip_md(row[j]) if j < len(row) else ""
                tbl.rows[i].cells[j].text = _xml_safe(val)
        doc.add_paragraph()

    i = 0
    n = len(lines)

    # ── H1 cover ──
    h1_title = title
    while i < n:
        stripped = lines[i].strip()
        m = _HEADING_RE.match(stripped)
        if m and len(m.group(1)) == 1:
            h1_title = _strip_md(m.group(2).strip())
            i += 1
            break
        i += 1
    if not h1_title:
        h1_title = md_path.stem.replace("_", " ").title()

    doc.add_paragraph()
    t = doc.add_heading(_xml_safe(h1_title), 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if case_id:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sub.add_run(f"Case: {case_id}  |  Generated {date_str}")
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    doc.add_paragraph()

    # Header table (preamble before the first ## heading)
    while i < n:
        stripped = lines[i].strip()
        if _HEADING_RE.match(stripped) and len(_HEADING_RE.match(stripped).group(1)) >= 2:
            break
        if _TABLE_ROW_RE.match(stripped):
            table_lines = []
            while i < n and _TABLE_ROW_RE.match(lines[i].strip()):
                table_lines.append(lines[i])
                i += 1
            _table(_parse_table_rows(table_lines))
            continue
        i += 1

    conf = doc.add_paragraph("CONFIDENTIAL — FOR AUTHORISED PERSONNEL ONLY")
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf.runs[0].font.bold = True
    conf.runs[0].font.color.rgb = RGBColor(0xef, 0x44, 0x44)
    doc.add_page_break()

    # ── Body ──
    in_blockquote = False
    while i < n:
        line = lines[i]
        stripped = line.strip()
        i += 1

        if not stripped:
            in_blockquote = False
            continue

        if stripped.startswith(">"):
            in_blockquote = True
            continue
        if in_blockquote:
            # Treat any immediately-following indented continuation of a
            # blockquote as still part of it; a normal line ends it.
            in_blockquote = False

        m = _HEADING_RE.match(stripped)
        if m:
            level = len(m.group(1))
            _h(_strip_md(m.group(2).strip()), min(level - 1, 8) if level > 1 else 1)
            continue

        if _TABLE_ROW_RE.match(stripped):
            table_lines = [line]
            while i < n and _TABLE_ROW_RE.match(lines[i].strip()):
                table_lines.append(lines[i])
                i += 1
            _table(_parse_table_rows(table_lines))
            continue

        if stripped == "---":
            continue

        bm = _BULLET_RE.match(stripped)
        if bm:
            _p(_strip_md(bm.group(1)), style="List Bullet")
            continue

        nm = _NUMBERED_RE.match(stripped)
        if nm:
            _p(_strip_md(nm.group(1)), style="List Number")
            continue

        _p(_strip_md(stripped))

    doc.save(str(out))
    print(f"[md_to_docx] DOCX saved: {out}")
    return out


# ── CLI ───────────────────────────────────────────────────────────────────

def _build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="Render a Markdown campaign report as a DOCX")
    p.add_argument("markdown", help="Path to input .md file")
    p.add_argument("--output", "-o", required=True, metavar="DOCX", help="Output DOCX path")
    p.add_argument("--case-id", metavar="ID", default="", help="Case ID shown under the title")
    p.add_argument("--title", metavar="TITLE", default="", help="Cover title (default: H1 of the document)")
    p.add_argument("--date", metavar="YYYY-MM-DD HH:MM UTC", default="", help="Report date (default: now UTC)")
    return p


if __name__ == "__main__":
    args = _build_parser().parse_args()
    convert(
        md_path=Path(args.markdown),
        output_path=Path(args.output),
        case_id=args.case_id,
        title=args.title,
        date_str=args.date,
    )
